import requests
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from enum import Enum
import time

current_time = time.time()
local_time = time.localtime(current_time)
formatted_time = time.strftime("%Y-%m-%d %H:%M:%S", local_time)

class MyParagraphAlignment(Enum):  # Unique name for the enum
  LEFT = 0
  CENTER = 1
  RIGHT = 2
  JUSTIFIED = 3


# The main website URL
main_url = 'https://www.youm7.com'





def get_article_links(url):
  """
  Fetches article links from the given URL and saves them to a file.

  Args:
      url (str): The URL of the webpage to scrape.

  Returns:
      list: A list of extracted article links (if any).
  """

  response = requests.get(url)

  # Check if the request was successful
  if response.status_code != 200:
    print(f"Failed to retrieve the main page from {url}")
    return []

  # Parse the HTML content using Beautiful Soup
  soup = BeautifulSoup(response.content, 'html.parser')

  # Find all divs with the class 'col-xs-4 smallSections'
  h2_element = soup.find('a', string=re.compile("أخبار عاجلة")).parent
  next_divs = h2_element.find_next_siblings("div")

  if h2_element:
    all_hrefs = []
    showMore_link = ""
    for div in next_divs:
      a_tags = div.find_all('a')
      for a_tag in a_tags:
        href = a_tag.get("href")  # Extract the href attribute
        if href:  # Check if href exists before appending
          href = href if href.startswith('http') else main_url + href
          if "Section" in href :   # separate the show more link
            showMore_link = href
            break
          all_hrefs.append(href)
    # Remove duplicates using set and convert back to list
    all_hrefs = list(set(all_hrefs)) 
  

    # Write extracted hrefs to a file 
    if all_hrefs:
      return showMore_link, all_hrefs  # Return the list of links 
        
    else:
      print("No hrefs found within anchor tags in the next divs.")
  else:
    print("h2 element with text 'أخبار عاجلة' not found.")

  return []  # Return an empty list if no links are found



def scrape_article(url):
  response = requests.get(url)
  
  # Check if the request was successful
  if response.status_code != 200:
    print(f"Failed to retrieve the article from {url}")
    return None
  
  
  # Parse the HTML content using Beautiful Soup
  soup = BeautifulSoup(response.content, 'html.parser')
  
  # Find the article content
  article_title = soup.find('h1')
  article_date = soup.find('span', class_='newsStoryDate')
  article_body = soup.find('div', id='articleBody')
  
  if article_title:
    # Extract the text content
    return article_title.get_text(strip=True), article_date.get_text(strip=True) ,article_body.get_text(strip=True)
  else:
    print(f"Article content not found on {url}")
    return None




def create_docx_report(articles):
  """
  Creates a docx report with formatted articles on separate pages, with right-aligned text.

  Args:
      articles (list): A list of tuples containing (article_title, article_body).
  """

  document = Document()

  for i, (title, date, body) in enumerate(articles):
    
    if i == 0:
      paragraph = document.add_paragraph()
    else:
      # Create a new section for subsequent articles
      new_section = document.add_section()

    # Add title paragraph
    paragraph = document.add_paragraph()
    paragraph.alignment = MyParagraphAlignment.RIGHT.value  
    title_run = paragraph.add_run(title)
    title_run.font.size = Pt(20)  # Set font size to 20pt
    title_run.font.name = 'Calibri'  
    paragraph.style = document.styles['Heading 1']  
    
    # Add data paragraph
    paragraph = document.add_paragraph()
    paragraph.alignment = MyParagraphAlignment.RIGHT.value  
    date_run = paragraph.add_run(date)
    date_run.font.size = Pt(10)  
    date_run.italic = True 
    date_run.font.name = 'Calibri'  

    # Add body paragraph
    paragraph = document.add_paragraph()
    paragraph.alignment = MyParagraphAlignment.RIGHT.value
    body_run = paragraph.add_run(body)
    body_run.italic = True  
    body_run.font.size = Pt(14)  

  # Save the document
  document.save(f'{formatted_time}_articles.docx')




showMoreLink, article_links = get_article_links(main_url)


# put the articles inside a .txt file for testing
# with open('article_urls.txt', 'w') as file:
#   for link in article_links:
#     article_title, article_body = scrape_article(link)
#     file.write(f"title: {article_title}\n")
#     file.write(f"body: {article_body}\n")
#     file.write('-' * 80 + '\n')

# get the article titles and body
articles = []
for link in article_links : 
  articles.append(scrape_article(link))



create_docx_report(articles)
